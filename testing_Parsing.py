from docx import *
import re

file = Document('path/syllabus.docx')
headFound = False
header = ""
data = ""

coreFile = file.core_properties


for par in file.paragraphs:
    
    for wInd, words in enumerate(par.runs):
        if words.bold:
            if re.match(r'^\s*$', words.text) or words.text == ':':
                exit
            elif wInd > 0:
                if not par.runs[wInd - 1].bold:
                    exit
                else:
                    if headFound == True and data != '':
                        data = ' '.join(data.split())
                        #Removes spaces between numbers
                        data = re.sub(r'(\d)\s+(\d)', r'\1\2', data)
                        #Removes spaces between colon and numbers
                        data = re.sub('(\d)\s[:]\s(\d)', r'\1:\2', data)
                        while(header[-1:] == ':' or header[-1:] == ' '):
                            header = header[:-1]
                        data = data.replace('%NL%', '\n')
                        print (header + ': ' + data + '\n')
                        header = ""
                    temp = words.text
                    while(temp[-1:] == ':' or temp[-1:] == ' '):
                        temp = temp[:-1]
                    while(temp[:1] == ':' or temp[:1] == ' '):
                        temp = temp[1:]
                    #print ("HEADER: " + temp + " CURRENT HEAD: " + header)
                    header = header + temp
                    headFound = True
                    newHead = True
                    data = ""
            else:
                if headFound == True and data != '':
                    data = ' '.join(data.split())
                    #Removes spaces between numbers
                    data = re.sub(r'(\d)\s+(\d)', r'\1\2', data)
                    #Removes spaces between colon and numbers
                    data = re.sub('(\d)\s[:]\s(\d)', r'\1:\2', data)
                    while(header[-1:] == ':' or header[-1:] == ' '):
                        header = header[:-1]
                    data = data.replace('%NL%', '\n')
                    print (header + ': ' + data + '\n')
                    header = ""
                temp = words.text
                while(temp[-1:] == ':' or temp[-1:] == ' '):
                    temp = temp[:-1]
                while(temp[:1] == ':' or temp[:1] == ' '):
                    temp = temp[1:]
                #print ("HEADER2: " + temp + " CURRENT HEAD2: " + header)
                header = header + temp
                headFound = True
                newHead = True
                data = ""
                
    if par.style.name=='Heading 1' or par.style.name=='Heading 2' or par.style.name=='Heading 3':
        if re.match(r'^\s*$', par.text) or par.text == ':':
            exit
        else:
            if headFound == True and data != '':
                data = ' '.join(data.split())
                #Removes spaces between numbers
                data = re.sub(r'(\d)\s+(\d)', r'\1\2', data)
                #Removes spaces between colon and numbers
                data = re.sub('(\d)\s[:]\s(\d)', r'\1:\2', data)
                while(header[-1:] == ':' or header[-1:] == ' '):
                    header = header[:-1]
                data = data.replace('%NL%', '\n')
                print (header + ': ' + data + '\n')
            temp = par.text
            while(temp[-1:] == ':' or temp[-1:] == ' '):
                temp = temp[:-1]
            while(temp[:1] == ':' or temp[:1] == ' '):
                temp = temp[1:]
            header = par.text
            headFound = True
            newHead = True
            data = ""
                
    if headFound == True:
        if data != "" and not re.match(r'^\s*$', par.text.strip()):
            data = data + '%NL%'
        if re.match(r'^\s*$', par.text.strip()) or words.text == ':' or par.style.name=='Heading 1' or par.style.name=='Heading 2' or par.style.name=='Heading 3':
            exit
        else:
            for wInd, words in enumerate(par.runs):
                if words.bold and (words.text in header or words.text in header + ':' or header in words.text):
                    exit
                else:
                    if data == '':
                        data = '%NL%'
                    temp = words.text
                    while(temp[:1] == ':'):
                        temp = temp[1:]
                    #if wInd == 0:
                    #    data = data + header + ' ' + temp
                    else:
                        data = data + temp

##for par in file.paragraphs:
##    print(par.text + 'SPLIT')

##for content in file.paragraphs:
##    if content.style.name=='Heading 1' or content.style.name=='Heading 2' or content.style.name=='Heading 3':
##        print (content.text)
##
##print(coreFile.title)

##for par in file.paragraphs:
##    if 'hyperlink' in par._element.getchildren():
##        par = par._element.getchildren().remove()
##        print("FIXED: " + par.text)
##    else:
##        print(par._element.getchildren())


##tables = file.tables
##for table in tables:
##    for row in table.rows:
##        for cell in row.cells:
##            for paragraph in cell.paragraphs:
##                print(paragraph.text)


##for par in file.paragraphs:
##    print(par.text)
##    break
    
        
